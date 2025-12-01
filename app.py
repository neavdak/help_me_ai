import streamlit as st
import torch
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from peft import PeftModel
import pypdf
import docx
import io
import os
import subprocess
import gc

def clear_gpu_memory():
    torch.cuda.empty_cache()
    gc.collect()

# --- CONFIGURATION ---
# --- CONFIGURATION ---
BASE_MODEL_ID = "Qwen/Qwen2.5-7B-Instruct"
# ADAPTER_PATH = "./final_model_qlora" # Disabled

# Set page title
st.set_page_config(page_title="Qwen 2.5 (7B) Chat", page_icon="🤖")
st.title("🤖 Qwen 2.5 (7B) - Production Ready")

# --- MODEL LOADING ---
@st.cache_resource
def load_model():
    """Loads the model and tokenizer. Cached to run only once."""
    print("Loading model...")
    bnb_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_quant_type="nf4",
        bnb_4bit_compute_dtype=torch.bfloat16
    )
    
    model = AutoModelForCausalLM.from_pretrained(
        BASE_MODEL_ID,
        quantization_config=bnb_config,
        device_map="auto"
    )
    
    # model = PeftModel.from_pretrained(base_model, ADAPTER_PATH) # Disabled
    tokenizer = AutoTokenizer.from_pretrained(BASE_MODEL_ID)
    return model, tokenizer

# Load model (this will show a spinner on first load)
with st.spinner("Loading model... This might take a minute."):
    model, tokenizer = load_model()

# --- SIDEBAR & CONFIG ---
with st.sidebar:
    st.header("Mode")
    app_mode = st.radio("Select Mode", ["Chat", "Doc Generator"])
    
    st.header("Parameters")
    max_new_tokens = st.slider("Max Length", min_value=64, max_value=4096, value=1024, step=64)
    temperature = st.slider("Creativity", min_value=0.1, max_value=1.5, value=0.7, step=0.1)

# --- CHAT MODE ---
if app_mode == "Chat":
    st.header("💬 Chat with Context")
    
    with st.sidebar:
        st.header("Upload Context")
        uploaded_file = st.file_uploader("Upload a file", type=["txt", "md", "py", "csv", "pdf", "docx"])
        
        file_context = ""
        if uploaded_file is not None:
            try:
                string_data = ""
                if uploaded_file.name.endswith(".pdf"):
                    pdf_reader = pypdf.PdfReader(uploaded_file)
                    for page in pdf_reader.pages:
                        string_data += page.extract_text() + "\n"
                elif uploaded_file.name.endswith(".docx"):
                    doc = docx.Document(uploaded_file)
                    for para in doc.paragraphs:
                        string_data += para.text + "\n"
                else:
                    # Text based files
                    string_data = uploaded_file.getvalue().decode("utf-8")
                
                file_context = f"Context from uploaded file ({uploaded_file.name}):\n{string_data}\n\n"
                st.success(f"Loaded {uploaded_file.name}!")
            except Exception as e:
                st.error(f"Error reading file: {e}")

    # Initialize chat history in session state
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Display chat messages from history on app rerun
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Accept user input
    if prompt := st.chat_input("What is up?"):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        # Display user message in chat message container
        with st.chat_message("user"):
            st.markdown(prompt)

        # Display assistant response in chat message container
        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            
            # Format prompt for DeepSeek (it handles raw text well, or we can use chat template)
            # DeepSeek R1 works best with a simple prompt or chat template. 
            # Let's use the chat template if available, otherwise raw text.
            
            messages = [{"role": "user", "content": f"{file_context}{prompt}"}]
            inputs = tokenizer.apply_chat_template(messages, return_tensors="pt", add_generation_prompt=True).to("cuda")
            
            clear_gpu_memory()
            with torch.no_grad():
                outputs = model.generate(
                    inputs, 
                    max_new_tokens=max_new_tokens,
                    do_sample=True, 
                    temperature=temperature,
                    top_p=0.9,
                )
                
            decoded_output = tokenizer.decode(outputs[0][len(inputs[0]):], skip_special_tokens=True)
            
            # Display response
            st.markdown(decoded_output)
            st.session_state.messages.append({"role": "assistant", "content": decoded_output})

# --- DOC GENERATOR MODE ---
elif app_mode == "Doc Generator":
    st.header("📄 Automated Documentation Generator")
    st.markdown("Generates documentation based on code changes and a template.")
    
    repo_path = st.text_input("Path to Code Folder (Git Repo)", value=os.getcwd())
    template_file = st.file_uploader("Upload Template (PDF/Word/Text)", type=["pdf", "docx", "txt", "md"])
    
    if st.button("Generate Documentation"):
        if not template_file:
            st.error("Please upload a template file.")
        else:
            with st.spinner("Analyzing code changes..."):
                # 1. Get Code Changes (Git Diff)
                diff_output = ""
                try:
                    # Try getting diff from last commit
                    result = subprocess.run(
                        ["git", "diff", "HEAD~1", "HEAD"], 
                        cwd=repo_path, 
                        capture_output=True, 
                        text=True,
                        check=True
                    )
                    diff_output = result.stdout
                    if not diff_output.strip():
                        # If no diff (maybe new repo), try getting all files
                        st.warning("No git diff found (or empty). Reading all .py files instead...")
                        for root, dirs, files in os.walk(repo_path):
                            if ".venv" in root or ".git" in root: continue
                            for file in files:
                                if file.endswith(".py"):
                                    path = os.path.join(root, file)
                                    with open(path, "r") as f:
                                        diff_output += f"\n--- File: {file} ---\n{f.read()}\n"
                except Exception as e:
                    st.error(f"Git error: {e}. Reading all .py files instead...")
                    # Fallback: Read all .py files
                    for root, dirs, files in os.walk(repo_path):
                        if ".venv" in root or ".git" in root: continue
                        for file in files:
                            if file.endswith(".py"):
                                path = os.path.join(root, file)
                                with open(path, "r") as f:
                                    diff_output += f"\n--- File: {file} ---\n{f.read()}\n"

                # 2. Read Template
                template_text = ""
                try:
                    if template_file.name.endswith(".pdf"):
                        pdf_reader = pypdf.PdfReader(template_file)
                        for page in pdf_reader.pages:
                            template_text += page.extract_text() + "\n"
                    elif template_file.name.endswith(".docx"):
                        doc = docx.Document(template_file)
                        for para in doc.paragraphs:
                            template_text += para.text + "\n"
                    else:
                        template_text = template_file.getvalue().decode("utf-8")
                except Exception as e:
                    st.error(f"Error reading template: {e}")

                # 3. Generate
                if diff_output and template_text:
                    prompt = f"""
                    You are a technical writer. Your task is to fill out the following documentation template based on the provided code changes.
                    
                    ### Template Structure:
                    {template_text}
                    
                    ### Code Changes/Context:
                    {diff_output[:10000]} 
                    (Truncated if too long)

                    ### Instructions:
                    - Analyze the code changes to understand what happened, where (file/line), and why.
                    - Fill out the template sections with this information.
                    - If a section is not relevant to the changes, write "No changes in this section."
                    - DO NOT ask clarifying questions.
                    - DO NOT explain what you are doing.
                    - DIRECTLY output the filled-out documentation in Markdown format.
                    """
                    
                    messages = [{"role": "user", "content": prompt}]
                    inputs = tokenizer.apply_chat_template(messages, return_tensors="pt", add_generation_prompt=True).to("cuda")
                    
                    clear_gpu_memory()
                    with torch.no_grad():
                        outputs = model.generate(
                            inputs, 
                            max_new_tokens=max_new_tokens,
                            do_sample=True, 
                            temperature=temperature,
                            top_p=0.9,
                        )
                    
                    decoded_output = tokenizer.decode(outputs[0][len(inputs[0]):], skip_special_tokens=True)
                    
                    final_report = decoded_output

                    st.markdown("### Generated Documentation")
                    st.markdown(final_report)
                    
                    # Download Button
                    st.download_button(
                        label="Download Report",
                        data=final_report,
                        file_name="generated_documentation.md",
                        mime="text/markdown"
                    )
                else:
                    st.error("Could not get code changes or template text.")
